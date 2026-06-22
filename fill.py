#!/usr/bin/env python3
"""
替换 Word 模板中的 {{变量}} 占位符
用法：python fill.py
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
import yaml
import os

# ============================================
# 1. 读取配置
# ============================================

with open("info.yaml", "r", encoding="utf-8") as f:
    INFO = yaml.safe_load(f)

print("=" * 50)
print("  读取配置")
print("=" * 50)
for k, v in INFO.items():
    print(f"  {{  {k}  }}  →  {v}")
print()

# ============================================
# 2. 替换函数
# ============================================

def replace_in_run(run, info):
    """替换单个 run 内的占位符"""
    for key, value in info.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, str(value))


def replace_in_paragraph(paragraph, info):
    """替换段落中的占位符，保留格式"""

    # 先尝试逐个 run 替换
    for run in paragraph.runs:
        replace_in_run(run, info)

    # 兜底：占位符跨多个 run 散落的情况
    full = paragraph.text
    replaced = full
    for key, value in info.items():
        replaced = replaced.replace(f"{{{{{key}}}}}", str(value))

    if replaced != full:
        # 有跨 run 的占位符，清空后重写
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = replaced


def replace_in_table(table, info):
    """替换表格中的所有占位符"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, info)


def replace_in_header_footer(doc, info):
    """替换页眉页脚中的占位符"""
    for section in doc.sections:
        if section.header and not section.header.is_linked_to_previous:
            for paragraph in section.header.paragraphs:
                replace_in_paragraph(paragraph, info)
            for table in section.header.tables:
                replace_in_table(table, info)

        if section.footer and not section.footer.is_linked_to_previous:
            for paragraph in section.footer.paragraphs:
                replace_in_paragraph(paragraph, info)
            for table in section.footer.tables:
                replace_in_table(table, info)


def fill(template_path, info, output_path):
    """主函数：读取模板 → 替换 → 保存"""
    print(f"📄 打开模板：{template_path}")
    doc = Document(template_path)

    # 段落
    para_count = 0
    for paragraph in doc.paragraphs:
        old = paragraph.text
        replace_in_paragraph(paragraph, info)
        if old != paragraph.text:
            para_count += 1
    print(f"   ✅ 替换了 {para_count} 个段落")

    # 表格
    cell_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    old = paragraph.text
                    replace_in_paragraph(paragraph, info)
                    if old != paragraph.text:
                        cell_count += 1
    print(f"   ✅ 替换了 {cell_count} 个表格单元格")

    # 页眉页脚
    replace_in_header_footer(doc, info)
    print(f"   ✅ 已处理页眉页脚")

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"\n🎉 输出：{output_path}")
    print(f"   用 Word 打开即可查看替换结果")


# ============================================
# 3. 运行
# ============================================

if __name__ == "__main__":
    fill("template.docx", INFO, "output/毕业论文_封面.docx")
